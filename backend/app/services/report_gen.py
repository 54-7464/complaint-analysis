from docx import Document
from docx.shared import Pt, Inches
import os


def extract_template_structure(file_path: str) -> str:
    """Extract structure and style from a Word template."""
    doc = Document(file_path)
    parts = []
    for p in doc.paragraphs:
        style = p.style.name if p.style else "Normal"
        text = p.text.strip()
        if text:
            parts.append(f"[{style}] {text}")
    for i, table in enumerate(doc.tables):
        parts.append(f"\n[表格{i+1}]")
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            parts.append(" | ".join(cells))
    return "\n".join(parts)


def generate_report_docx(template_path: str, analysis_summary: str,
                         report_content: str, output_path: str) -> str:
    """Generate a Word report based on template structure + AI-generated content."""
    doc = Document(template_path) if os.path.exists(template_path) else Document()

    # Replace or append content
    if doc.paragraphs:
        # Replace the first substantial paragraph
        for p in doc.paragraphs:
            if len(p.text) > 20:
                p.text = report_content
                break
        else:
            doc.add_paragraph(report_content)
    else:
        doc.add_heading("卫生投诉分析报告", level=0)
        doc.add_paragraph(report_content)

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    return output_path
