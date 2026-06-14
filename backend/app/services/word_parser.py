from docx import Document


def parse_word(file_path: str) -> str:
    """Extract all text from a Word document."""
    doc = Document(file_path)
    paragraphs = []
    for p in doc.paragraphs:
        if p.text.strip():
            paragraphs.append(p.text.strip())
    # also extract table text
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))
    return "\n".join(paragraphs)
