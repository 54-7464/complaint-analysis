import json
import os
from fastapi import APIRouter, Depends, File, HTTPException, UploadFile, Query
from fastapi.responses import FileResponse
from sqlalchemy.orm import Session
from pydantic import BaseModel
from jose import JWTError, jwt
from ..database import get_db
from ..models.user import User
from ..models.user_config import UserAIConfig
from ..models.project import Project
from ..models.labeling import LabelingJob, Label
from ..models.report import Report
from ..auth import get_current_user
from ..config import settings
from ..services.excel_handler import parse_excel
from ..services.ai_labeler import call_ai
from ..services.report_gen import extract_template_structure, generate_report_docx
from ..services.crypto import decrypt_api_key
from ..services.security import validate_upload_path

router = APIRouter(prefix="/api/report", tags=["report"])


class GenerateReportRequest(BaseModel):
    project_id: int
    labeling_job_id: int
    template_path: str


@router.post("/upload-template/{project_id}")
async def upload_template(project_id: int, file: UploadFile = File(...),
                          db: Session = Depends(get_db), user: User = Depends(get_current_user)):
    project = db.query(Project).filter(Project.id == project_id, Project.user_id == user.id).first()
    if not project:
        raise HTTPException(404, "项目不存在")
    ext = os.path.splitext(file.filename or ".docx")[1].lower()
    if ext not in {".docx"}:
        raise HTTPException(400, "仅支持 .docx 文件")
    user_dir = os.path.join(settings.UPLOAD_DIR, str(user.id), str(project_id))
    os.makedirs(user_dir, exist_ok=True)
    import uuid
    safe_name = f"template_{uuid.uuid4().hex}{ext}"
    save_path = os.path.join(user_dir, safe_name)
    content = await file.read()
    with open(save_path, "wb") as f:
        f.write(content)
    return {"file_path": save_path, "filename": file.filename}

@router.delete("/template/{project_id}")
def delete_template(project_id: int, path: str = Query(...),
                    db: Session = Depends(get_db), user: User = Depends(get_current_user)):
    """删除已上传的报告模板文件"""
    project = db.query(Project).filter(Project.id == project_id, Project.user_id == user.id).first()
    if not project:
        raise HTTPException(404, '项目不存在')
    safe_path = validate_upload_path(path)
    try:
        if os.path.exists(safe_path):
            os.remove(safe_path)
    except OSError:
        pass
    return {"ok": True}



@router.post("/generate")
def generate_report(req: GenerateReportRequest, db: Session = Depends(get_db),
                    user: User = Depends(get_current_user)):
    project = db.query(Project).filter(Project.id == req.project_id, Project.user_id == user.id).first()
    if not project:
        raise HTTPException(404, "项目不存在")
    job = db.query(LabelingJob).filter(LabelingJob.id == req.labeling_job_id).first()
    if not job:
        raise HTTPException(404, "标注任务不存在")
    cfg = db.query(UserAIConfig).filter(UserAIConfig.user_id == user.id).first()
    if not cfg or not cfg.api_key_encrypted:
        raise HTTPException(400, "请先配置 AI API Key")
    api_key = decrypt_api_key(cfg.api_key_encrypted)
    template_structure = extract_template_structure(req.template_path) if os.path.exists(req.template_path) else ""
    labels = db.query(Label).filter(Label.labeling_job_id == job.id).all()
    label_names = [l.name for l in labels]
    excel_path = job.labeled_file_path or job.data_source.file_path
    parsed = parse_excel(excel_path) if os.path.exists(excel_path) else {"rows": [], "columns": []}
    from collections import Counter
    label_counter = Counter()
    for row in parsed["rows"]:
        if row:
            labels_str = str(row[-1]) if row[-1] else ""
            for lbl in labels_str.split(","):
                lbl = lbl.strip()
                if lbl:
                    label_counter[lbl] += 1
    summary = f"""数据总量：{parsed['row_count']} 条, 标签种类：{len(label_names)} 个, 标签分布：{json.dumps(dict(label_counter), ensure_ascii=False, indent=2)}"""
    prompt = f"""你是专业的公共卫生数据分析师。根据以下数据撰写卫生投诉分析报告。
【报告模板结构参考】{template_structure[:3000]}
【数据分析摘要】{summary}
要求：仿照模板结构框架和语句风格，使用专业公文体语言，包含主要发现、数据解读和建议。"""
    try:
        ai_response = call_ai(prompt, api_key, cfg.base_url, cfg.model_name)
    except Exception as e:
        ai_response = f"AI 调用失败：{str(e)}\n\n请检查 API Key 和网络配置。"
    user_dir = os.path.join(settings.UPLOAD_DIR, str(user.id), str(req.project_id), "reports")
    os.makedirs(user_dir, exist_ok=True)
    import uuid
    output_path = os.path.join(user_dir, f"report_{uuid.uuid4().hex}.docx")
    if os.path.exists(req.template_path):
        generate_report_docx(req.template_path, summary, ai_response, output_path)
    else:
        from docx import Document
        doc = Document()
        doc.add_heading("卫生投诉分析报告", level=0)
        doc.add_paragraph(ai_response)
        doc.save(output_path)
    report = Report(project_id=req.project_id, template_path=req.template_path, content=ai_response, file_path=output_path)
    db.add(report)
    db.commit()
    db.refresh(report)
    return {"id": report.id, "content": ai_response, "file_path": output_path}


@router.get("/download/{report_id}")
def download_report(report_id: int, token: str = Query(None), db: Session = Depends(get_db)):
    if not token:
        raise HTTPException(401, "需要认证")
    try:
        payload = jwt.decode(token, settings.SECRET_KEY, algorithms=[settings.ALGORITHM])
        user_id = int(payload.get("sub"))
    except (JWTError, ValueError):
        raise HTTPException(401, "token 无效")
    report = db.query(Report).filter(Report.id == report_id).first()
    if not report:
        raise HTTPException(404, "报告不存在")
    project = db.query(Project).filter(Project.id == report.project_id, Project.user_id == user_id).first()
    if not project:
        raise HTTPException(404, "无权限")
    if not report.file_path or not os.path.exists(report.file_path):
        raise HTTPException(404, "报告文件不存在")
    return FileResponse(report.file_path, filename=f"分析报告_{report.id}.docx",
                        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")


# ══════════════════════════════════════════════════════
#  AI 自由表格分析
# ══════════════════════════════════════════════════════

class TableAnalysisRequest(BaseModel):
    table_text: str       # 粘贴的表格内容（CSV/TSV/Markdown）
    instruction: str      # 自然语言描述的分析需求
    model: str = ""       # 可选覆盖模型


@router.post("/analyze-table")
def analyze_table(req: TableAnalysisRequest, db: Session = Depends(get_db),
                  user: User = Depends(get_current_user)):
    """粘贴表格 + 分析指令 → AI 返回分析结果"""
    cfg = db.query(UserAIConfig).filter(UserAIConfig.user_id == user.id).first()
    if not cfg or not cfg.api_key_encrypted:
        raise HTTPException(400, "请先配置 AI API Key")

    api_key = decrypt_api_key(cfg.api_key_encrypted)
    model = req.model or cfg.model_name

    prompt = f"""你是一个专业的数据分析师。请分析以下表格数据，严格按用户指令输出结果。

【用户分析指令】
{req.instruction}

【表格数据】
{req.table_text}

【输出要求】
- 只输出分析结果，不要重复原始数据
- 使用 Markdown 格式：表格用 markdown table，重点用 **加粗**
- 如果有数值计算，写出计算过程和公式
- 给出清晰的结论或建议"""

    try:
        result = call_ai(prompt, api_key, cfg.base_url, model)
    except Exception as e:
        result = f"❌ AI 调用失败：{str(e)}"

    return {"result": result, "model_used": model}


class AnalyzeJobRequest(BaseModel):
    job_id: int
    instruction: str
    model: str = ''


@router.post('/analyze-job')
def analyze_job(req: AnalyzeJobRequest, db: Session = Depends(get_db),
                user: User = Depends(get_current_user)):
    """对流水线标注结果进行AI自然语言分析"""
    cfg = db.query(UserAIConfig).filter(UserAIConfig.user_id == user.id).first()
    if not cfg or not cfg.api_key_encrypted:
        raise HTTPException(400, "请先配置 AI API Key")

    job = db.query(LabelingJob).filter(LabelingJob.id == req.job_id).first()
    if not job:
        raise HTTPException(404, "标注任务不存在")
    project = db.query(Project).filter(Project.id == job.project_id, Project.user_id == user.id).first()
    if not project:
        raise HTTPException(404, "无权限")

    api_key = decrypt_api_key(cfg.api_key_encrypted)
    model = req.model or cfg.model_name

    labels = db.query(Label).filter(Label.labeling_job_id == job.id).all()
    label_names = [l.name for l in labels]

    excel_path = job.labeled_file_path or job.data_source.file_path
    summary_text = ""
    if excel_path and os.path.exists(excel_path):
        parsed = parse_excel(excel_path)
        from collections import Counter
        label_counter = Counter()
        lci = _find_label_column_fast(parsed["columns"])
        for row in parsed["rows"]:
            if lci is not None and lci < len(row) and row[lci]:
                for lbl in str(row[lci]).replace("，", ",").split(","):
                    lbl = lbl.strip()
                    if lbl:
                        label_counter[lbl] += 1
        summary_text = f'数据总量: {parsed.row_count} 条\\n标签种类: {len(label_names)} 个\\n标签分布:\\n{json.dumps(dict(label_counter), ensure_ascii=False, indent=2)}'

    prompt = f'''你是专业数据分析师。基于以下标注结果进行分析。\\n\\n【分析指令】\\n{req.instruction}\\n\\n【标注数据摘要】\\n{summary_text}\\n\\n【输出要求】\\n- 直接回答分析指令\\n- 使用 Markdown 格式\\n- 如有数值计算写出过程和公式'''

    try:
        result = call_ai(prompt, api_key, cfg.base_url, model)
    except Exception as e:
        result = f"AI调用失败: {str(e)}"

    return {"result": result, "model_used": model}


def _find_label_column_fast(columns: list) -> int | None:
    for j, col in enumerate(columns):
        if ("标签" in str(col).strip() and "思考" not in str(col).strip()):
            return j
    if len(columns) >= 2:
        return len(columns) - 2
    return None
